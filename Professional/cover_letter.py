from docx import Document
from datetime import datetime
import pyttsx3

# Initialize the text-to-speech engine
engine = pyttsx3.init()

def speak(text):
    engine.say(text)
    engine.runAndWait()

def get_user_input(prompt):
    speak(prompt)  # Speak the prompt
    return input(prompt)

def add_heading_and_paragraph(document, heading, content):
    document.add_heading(heading, level=2)
    document.add_paragraph(content)
    speak(f"{heading}: {content}")

def create_cover_letter():
    speak("Sure. Launching the automated cover letter creation program. Please provide the following information. You can always press Enter to skip.")

    user_input_data = {
        'sender_name': get_user_input("Enter your full name: "),
        'sender_address': get_user_input("Enter your address: "),
        'current_date': datetime.now().strftime('%Y-%m-%d'),
        'recipient_name': get_user_input("Enter recipient's full name: "),
        'company_name': get_user_input("Enter company name: "),
        'position': get_user_input("Enter the position you are applying for: "),
        'intro_paragraph': get_user_input("Write an introduction paragraph: "),
        'main_paragraph': get_user_input("Write the main content of your letter: "),
        'closing_paragraph': get_user_input("Write a closing paragraph: "),
    }

    document = Document()

    # Print user_data for debugging
    print("Debugging user_data:", user_input_data)

    # Add header with current date
    header = f"{user_input_data['sender_name']} - Cover Letter"
    document.add_heading(header, level=1)
    document.add_paragraph(f"Date: {user_input_data['current_date']}")
    speak(f"Creating cover letter for {user_input_data['sender_name']}")

    # Add sender and recipient information
    add_heading_and_paragraph(document, 'Sender Information', f"Name: {user_input_data['sender_name']}\nAddress: {user_input_data['sender_address']}")
    add_heading_and_paragraph(document, 'Recipient Information', f"Name: {user_input_data['recipient_name']}\nCompany: {user_input_data['company_name']}")

    # Add position and introduction
    add_heading_and_paragraph(document, 'Position Applied For', user_input_data['position'])
    add_heading_and_paragraph(document, 'Introduction', user_input_data['intro_paragraph'])

    # Add main content
    add_heading_and_paragraph(document, 'Main Content', user_input_data['main_paragraph'])

    # Add closing
    add_heading_and_paragraph(document, 'Closing', user_input_data['closing_paragraph'])

    # Save the document
    document.save(f'{user_input_data["sender_name"]}_cover_letter.docx')
    speak(f"Cover letter for {user_input_data['sender_name']} created successfully. Saved as {user_input_data['sender_name']}_cover_letter.docx")

if __name__ == "__main__":
    create_cover_letter()
