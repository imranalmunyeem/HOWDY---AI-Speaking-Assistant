from docx import Document
from datetime import datetime
import pyttsx3

# Initialize the text-to-speech engine
engine = pyttsx3.init()

def speak(text):
    engine.say(text)
    engine.runAndWait()

def get_user_input(prompt):
    speak(prompt)  # Speak the prompt
    return input(prompt)

def add_heading_and_paragraph(document, heading, content):
    document.add_heading(heading, level=2)
    document.add_paragraph(content)
    speak(f"{heading}: {content}")

def create_resume():
    speak("Sure. Launching the automated resume creation program. Please provide the following information. You can always press Enter to skip.")

    user_input_data = {
        'name': get_user_input("Enter your full name: "),
        'address': get_user_input("Enter your address: "),
        'phone': get_user_input("Enter your phone number: "),
        'email': get_user_input("Enter your email address: "),
        'skills': get_user_input("Enter your skills (comma-separated): ").split(','),
        'experience': get_user_input("Enter your work experience: "),
        'education': get_user_input("Enter your education background: "),
        'projects': get_user_input("Enter your projects: "),
        'certifications': get_user_input("Enter your certifications: "),
    }

    document = Document()

    # Print user_data for debugging
    print("Debugging user_data:", user_input_data)

    # Add header with current date
    header = f"{user_input_data['name']} - Resume"
    document.add_heading(header, level=1)
    document.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
    speak(f"Creating resume for {user_input_data['name']}")

    # Add contact information if provided
    contact_info = ""
    if user_input_data['address'] or user_input_data['phone'] or user_input_data['email']:
        add_heading_and_paragraph(document, 'Contact Information', f"Address: {user_input_data['address']}\nPhone: {user_input_data['phone']}\nEmail: {user_input_data['email']}")
        contact_info = f"Address: {user_input_data['address']}, Phone: {user_input_data['phone']}, Email: {user_input_data['email']}"

    # Add skills if provided
    if user_input_data['skills']:
        add_heading_and_paragraph(document, 'Skills', ', '.join(user_input_data['skills']))

    # Add work experience if provided
    if user_input_data['experience']:
        add_heading_and_paragraph(document, 'Work Experience', user_input_data['experience'])

    # Add education if provided
    if user_input_data['education']:
        add_heading_and_paragraph(document, 'Education', user_input_data['education'])

    # Add projects if provided
    if user_input_data['projects']:
        add_heading_and_paragraph(document, 'Projects', user_input_data['projects'])

    # Add certifications if provided
    if user_input_data['certifications']:
        add_heading_and_paragraph(document, 'Certifications', user_input_data['certifications'])

    # Save the document
    document.save(f'{user_input_data["name"]}_resume.docx')
    speak(f"Resume for {user_input_data['name']} created successfully. Saved as {user_input_data['name']}_resume.docx")

if __name__ == "__main__":
    create_resume()
